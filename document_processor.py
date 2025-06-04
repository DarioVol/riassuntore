"""
Modulo per l'estrazione e processamento di testi da diversi formati di file.

Supporta: PDF, DOCX, DOC, TXT
"""

import os
from pathlib import Path
from typing import Optional
import PyPDF2
import docx


class DocumentProcessor:
    """
    Processore di documenti per diversi formati di file.

    Supporta estrazione di testo da:
    - PDF (tramite PyPDF2)
    - DOCX/DOC (tramite python-docx)
    - TXT (lettura diretta)
    """

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}

    def __init__(self):
        """Inizializza il processore di documenti."""
        self.last_processed_file = None
        self.last_extracted_text = None

    def is_supported_format(self, file_path: str) -> bool:
        """
        Verifica se il formato del file è supportato.

        Args:
            file_path (str): Percorso del file da verificare

        Returns:
            bool: True se supportato, False altrimenti
        """
        extension = Path(file_path).suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS

    def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Estrae testo da file PDF usando PyPDF2.

        Args:
            file_path (str): Percorso del file PDF

        Returns:
            str: Testo estratto

        Raises:
            FileNotFoundError: Se il file non esiste
            Exception: Per errori di lettura PDF
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File PDF non trovato: {file_path}")

        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Solo se la pagina contiene testo
                            text_content.append(page_text)
                    except Exception as e:
                        print(f"Warning: Errore nella pagina {page_num + 1}: {e}")
                        continue

            extracted_text = "\n".join(text_content).strip()

            if not extracted_text:
                raise Exception("PDF non contiene testo estraibile o è protetto")

            return extracted_text

        except Exception as e:
            raise Exception(f"Errore nell'estrazione da PDF: {str(e)}")

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Estrae testo da file DOCX/DOC usando python-docx.

        Args:
            file_path (str): Percorso del file DOCX/DOC

        Returns:
            str: Testo estratto

        Raises:
            FileNotFoundError: Se il file non esiste
            Exception: Per errori di lettura DOCX
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File DOCX non trovato: {file_path}")

        try:
            doc = docx.Document(file_path)
            text_content = []

            # Estrae testo dai paragrafi
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Estrae testo dalle tabelle
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())

            extracted_text = "\n".join(text_content).strip()

            if not extracted_text:
                raise Exception("File DOCX vuoto o non contiene testo")

            return extracted_text

        except Exception as e:
            raise Exception(f"Errore nell'estrazione da DOCX: {str(e)}")

    def extract_text_from_txt(self, file_path: str) -> str:
        """
        Estrae testo da file TXT con encoding automatico.

        Args:
            file_path (str): Percorso del file TXT

        Returns:
            str: Testo estratto

        Raises:
            FileNotFoundError: Se il file non esiste
            Exception: Per errori di lettura TXT
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File TXT non trovato: {file_path}")

        # Prova diversi encoding
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read().strip()

                if content:
                    return content
                else:
                    raise Exception("File TXT vuoto")

            except UnicodeDecodeError:
                continue
            except Exception as e:
                raise Exception(f"Errore nella lettura del file TXT: {str(e)}")

        raise Exception(f"Impossibile decodificare il file TXT con encoding supportati")

    def extract_text(self, file_path: str) -> str:
        """
        Router principale per estrazione testo basato sull'estensione del file.

        Args:
            file_path (str): Percorso del file da processare

        Returns:
            str: Testo estratto dal file

        Raises:
            ValueError: Se formato non supportato
            FileNotFoundError: Se file non trovato
            Exception: Per errori di estrazione
        """
        file_path = str(file_path)  # Assicura che sia stringa

        if not self.is_supported_format(file_path):
            extension = Path(file_path).suffix.lower()
            raise ValueError(
                f"Formato file '{extension}' non supportato. "
                f"Formati supportati: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )

        # Cache dell'ultimo file processato
        if self.last_processed_file == file_path and self.last_extracted_text:
            return self.last_extracted_text

        file_extension = Path(file_path).suffix.lower()

        # Router per tipo di file
        if file_extension == '.pdf':
            extracted_text = self.extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            extracted_text = self.extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            extracted_text = self.extract_text_from_txt(file_path)
        else:
            # Questo non dovrebbe mai accadere data la verifica precedente
            raise ValueError(f"Estensione non gestita: {file_extension}")

        # Salva in cache
        self.last_processed_file = file_path
        self.last_extracted_text = extracted_text

        return extracted_text

    def get_file_info(self, file_path: str) -> dict:
        """
        Ottiene informazioni sul file.

        Args:
            file_path (str): Percorso del file

        Returns:
            dict: Informazioni sul file (nome, estensione, dimensione, supportato)
        """
        path_obj = Path(file_path)

        return {
            'filename': path_obj.name,
            'extension': path_obj.suffix.lower(),
            'size_bytes': path_obj.stat().st_size if path_obj.exists() else 0,
            'size_mb': round(path_obj.stat().st_size / (1024*1024), 2) if path_obj.exists() else 0,
            'exists': path_obj.exists(),
            'is_supported': self.is_supported_format(file_path)
        }